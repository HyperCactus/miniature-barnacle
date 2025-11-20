"""
Document parser module for extracting text from various document formats.
Supports PDF, DOCX, TXT, and MD files.
"""

import io
from pathlib import Path
from typing import Union
import PyPDF2
from docx import Document
import markdown


def parse_pdf(file_obj) -> str:
    """Extract text from PDF file."""
    text_content = []
    
    try:
        pdf_reader = PyPDF2.PdfReader(file_obj)
        num_pages = len(pdf_reader.pages)
        
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text = page.extract_text()
            if text:
                text_content.append(text)
        
        return "\n\n".join(text_content)
    except Exception as e:
        raise Exception(f"Error parsing PDF: {str(e)}")


def parse_docx(file_obj) -> str:
    """Extract text from DOCX file."""
    try:
        doc = Document(file_obj)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
        
        return "\n\n".join(text_content)
    except Exception as e:
        raise Exception(f"Error parsing DOCX: {str(e)}")


def parse_txt(file_obj) -> str:
    """Extract text from TXT file."""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                file_obj.seek(0)
                content = file_obj.read()
                if isinstance(content, bytes):
                    text = content.decode(encoding)
                else:
                    text = content
                return text
            except (UnicodeDecodeError, AttributeError):
                continue
        
        raise Exception("Could not decode text file with supported encodings")
    except Exception as e:
        raise Exception(f"Error parsing TXT: {str(e)}")


def parse_markdown(file_obj) -> str:
    """Extract text from Markdown file."""
    try:
        file_obj.seek(0)
        content = file_obj.read()
        
        if isinstance(content, bytes):
            content = content.decode('utf-8')
        
        # Convert markdown to plain text (remove markdown syntax)
        # We'll use the markdown library to convert to HTML, then strip tags
        html = markdown.markdown(content)
        
        # Simple HTML tag removal
        import re
        text = re.sub('<[^<]+?>', '', html)
        text = re.sub('&nbsp;', ' ', text)
        text = re.sub('&lt;', '<', text)
        text = re.sub('&gt;', '>', text)
        text = re.sub('&amp;', '&', text)
        
        return text
    except Exception as e:
        raise Exception(f"Error parsing Markdown: {str(e)}")


def parse_document(file_obj) -> str:
    """
    Parse document and extract text based on file type.
    
    Args:
        file_obj: Uploaded file object from Streamlit
        
    Returns:
        Extracted text as string
    """
    file_name = file_obj.name.lower()
    
    # Reset file pointer to beginning
    file_obj.seek(0)
    
    if file_name.endswith('.pdf'):
        return parse_pdf(file_obj)
    elif file_name.endswith('.docx'):
        return parse_docx(file_obj)
    elif file_name.endswith('.txt'):
        return parse_txt(file_obj)
    elif file_name.endswith('.md'):
        return parse_markdown(file_obj)
    else:
        raise ValueError(f"Unsupported file format: {file_name}")


def clean_text(text: str) -> str:
    """
    Clean extracted text for better TTS processing.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    import re
    
    # Remove excessive whitespace
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Remove page numbers and headers/footers patterns
    text = re.sub(r'\n\d+\n', '\n', text)
    
    # Fix common OCR errors and formatting issues
    text = re.sub(r'([a-z])- ([a-z])', r'\1\2', text)  # Fix hyphenated words
    
    return text.strip()
